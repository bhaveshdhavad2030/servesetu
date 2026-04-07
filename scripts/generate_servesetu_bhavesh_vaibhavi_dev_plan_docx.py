from __future__ import annotations

from datetime import date
from pathlib import Path


def _h(doc, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def _p(doc, text: str) -> None:
    doc.add_paragraph(text)


def _bullets(doc, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _table(doc, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build_doc(output_path: Path) -> None:
    from docx import Document

    doc = Document()

    customer_term = "Booking"
    technician_term = "Work Order"
    internal_term = "Service Request"

    # Title
    doc.add_heading("ServeSetu — Dev Work Plan (Bhavesh + Vaibhavi) + Branch Strategy", level=0)
    _p(doc, f"Generated on: {date.today().isoformat()}")
    _p(
        doc,
        "This document aligns development work between Bhavesh and Vaibhavi for a web-first MVP, "
        "and defines branch naming + merge workflow to keep tracking simple.",
    )
    _p(
        doc,
        f"Terminology: Customer='{customer_term}', Technician='{technician_term}', Internal='{internal_term}'.",
    )

    doc.add_page_break()

    # 1) Tech stack lock (current + decisions)
    _h(doc, "1) Technology lock (what we are using)", 1)
    _h(doc, "Frontend (current in repo)", 2)
    _bullets(
        doc,
        [
            "React 18 + TypeScript",
            "Vite (dev server/build)",
            "React Router",
            "Tailwind CSS",
            "Redux Toolkit (state management scaffold)",
            "Axios (API calls scaffold)",
        ],
    )

    _h(doc, "Backend + database (to implement next)", 2)
    _p(
        doc,
        "Note: backend folder is currently empty, so this is the recommended lock for web-first speed and simplicity.",
    )
    _bullets(
        doc,
        [
            "Backend: Node.js + Express + TypeScript",
            "Database: PostgreSQL",
            "ORM: Prisma",
            "Auth: OTP (SMS) + JWT/session; role-based access (Customer/Technician/Admin)",
            "File storage: S3-compatible (or local in dev) for before/after photos",
            "Payments: Razorpay (India) with webhook verification",
        ],
    )

    doc.add_page_break()

    # 2) What is done in frontend vs remaining
    _h(doc, "2) Current status vs remaining (as per blueprint)", 1)
    _h(doc, "Frontend already started (present UI pages)", 2)
    _bullets(
        doc,
        [
            "Landing page",
            "Marketplace listing (sample technicians/services data)",
            "Technician profile page (UI)",
            "Booking flow UI (multi-step form) with static fees",
            "Customer dashboard UI (sample bookings)",
            "Technician dashboard UI (basic)",
            "Login placeholder page",
        ],
    )

    _h(doc, "Major gaps to match the ServeSetu blueprint", 2)
    _bullets(
        doc,
        [
            "No real backend APIs yet (everything is static data).",
            "No admin portal yet (required for verification, catalog/pricing, disputes, monitoring).",
            "No real booking/work-order status tracking (state machine) across customer/technician/admin.",
            "No estimate builder + customer approval (OTP/signature) flow.",
            "No payments integration + invoices + webhooks.",
            "No photo uploads (before/after) storage pipeline.",
            "No database schema or migrations.",
            "Service catalog data currently includes non-MVP services (plumbing/painting etc.) and needs alignment to your electronics categories.",
        ],
    )

    doc.add_page_break()

    # 3) Branch strategy
    _h(doc, "3) Branch strategy (easy tracking and merging)", 1)
    _p(doc, "Base branch: develop. Each feature/epic gets its own branch and PR into develop.")

    _h(doc, "Branch naming convention", 2)
    _bullets(
        doc,
        [
            "feat/<area>-<short-name>  (new feature)",
            "fix/<area>-<short-name>   (bug fix)",
            "chore/<area>-<short-name> (refactor, tooling, cleanup)",
        ],
    )

    _h(doc, "Recommended branches for MVP", 2)
    _bullets(
        doc,
        [
            "chore/backend-bootstrap (create backend app, linting, env, healthcheck)",
            "feat/auth-otp-roles (OTP login + roles)",
            "feat/catalog-services-pricing (service categories + issues + pricing rules)",
            "feat/bookings-core (customer booking create + list + details)",
            "feat/work-orders-core (technician accept/decline + status updates)",
            "feat/estimate-approval (estimate builder + customer approval + audit log)",
            "feat/payments-invoices (Razorpay + webhooks + invoices)",
            "feat/uploads-before-after (photo upload + linking to requests)",
            "feat/admin-portal (monitoring + verification + catalog editor + refunds)",
        ],
    )

    _h(doc, "Merge workflow", 2)
    _bullets(
        doc,
        [
            "Small PRs (1–3 days work) are preferred.",
            "PR must include: screenshots (frontend) or API examples (backend) + quick test notes.",
            "Never work directly on develop; always via feature branch.",
        ],
    )

    doc.add_page_break()

    # 4) Work split (Bhavesh vs Vaibhavi)
    _h(doc, "4) Work split (Bhavesh vs Vaibhavi)", 1)
    _p(doc, "Goal: parallelize backend foundations with frontend flow wiring, so both devs stay unblocked.")

    _h(doc, "Bhavesh (backend + platform owner)", 2)
    _bullets(
        doc,
        [
            "Backend bootstrap (Express + TS + Prisma + Postgres) + env config + healthcheck",
            "DB schema: users, technician profiles, service catalog, service requests, estimates, payments, reviews, audit logs",
            "Core APIs: auth, catalog, bookings/work orders, status transitions",
            "Admin portal backend endpoints (verification, refunds, reports)",
            "Payments integration (Razorpay) + webhook verification",
            "Deploy staging (backend + frontend) and keep it stable",
        ],
    )

    _h(doc, "Vaibhavi (frontend owner)", 2)
    _bullets(
        doc,
        [
            "Refine and align UI to ServeSetu electronics categories (AC/WM/TV/Fridge/Geyser/RO/Cooler/Speaker)",
            "Replace static sample data with API integration (Axios) + loading/error states",
            "Customer portal: marketplace → booking → tracking timeline → invoice → review",
            "Technician portal: work order list → accept/decline → status updates → diagnosis → estimate builder → photo upload",
            "Admin portal UI shell (tables/forms) once backend endpoints exist",
            "Consistent navigation, routes, and terminology (Booking / Work Order)",
        ],
    )

    _h(doc, "Shared work (pair decisions)", 2)
    _bullets(
        doc,
        [
            "Lock MVP city + 2–3 services for first pilot (then expand).",
            "Lock pricing rules (visit fee/platform fee/parts extra/cancellation/refund).",
            "Lock status machine transitions and SLA timers.",
        ],
    )

    doc.add_page_break()

    # 5) Week 1–2 immediate tasks (start now)
    _h(doc, "5) Immediate next tasks (start now)", 1)
    _h(doc, "Week 1 (foundation)", 2)
    _table(
        doc,
        headers=["Task", "Owner", "Output"],
        rows=[
            ["Decide MVP city + 2–3 service categories", "Bhavesh + Vaibhavi", "Final list in docs"],
            ["Finalize pricing rules + cancellation/refund rules", "Ops/Docs + Bhavesh", "Policy text + rules table"],
            ["Backend bootstrap + Postgres + Prisma schema v1", "Bhavesh", "Running API + migrations"],
            ["Frontend: align service catalog content to electronics", "Vaibhavi", "Updated UI + data models"],
            ["API contract v1 (endpoints + payloads)", "Bhavesh", "Shared spec for frontend"],
        ],
    )

    _h(doc, "Week 2 (first end-to-end demo)", 2)
    _table(
        doc,
        headers=["Task", "Owner", "Output"],
        rows=[
            ["OTP auth + roles", "Bhavesh", "Login works + role guards"],
            ["Customer creates booking (API + UI)", "Vaibhavi + Bhavesh", "Booking create/list/details"],
            ["Technician work order list + accept", "Vaibhavi + Bhavesh", "Accept/decline + status"],
            ["Admin basic monitor (read-only)", "Bhavesh", "Admin sees all requests"],
        ],
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    output_path = repo_root / "docs" / "ServeSetu_Bhavesh_Vaibhavi_Dev_Plan.docx"
    build_doc(output_path)
    print(f"Wrote: {output_path}")


if __name__ == "__main__":
    main()

