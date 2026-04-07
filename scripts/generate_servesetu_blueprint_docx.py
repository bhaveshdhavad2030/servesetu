from __future__ import annotations

from datetime import date
from pathlib import Path


def _add_heading(doc, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_bullets(doc, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_numbered(doc, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def _add_kv_table(doc, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Details"
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = k
        cells[1].text = v


def build_doc(output_path: Path) -> None:
    from docx import Document

    doc = Document()

    customer_term = "Booking"
    technician_term = "Work Order"
    internal_term = "Service Request"

    # Title
    doc.add_heading("ServeSetu — Product Blueprint (Customer + Technician + Admin)", level=0)
    doc.add_paragraph(f"Generated on: {date.today().isoformat()}")
    doc.add_paragraph(
        "This document converts the shared business plan into an implementation-ready blueprint: "
        "development path, app workflows, and app structure (pages, APIs, modules, data)."
    )
    doc.add_paragraph(
        f"Terminology used in this document: Customer sees '{customer_term}', Technician sees '{technician_term}', "
        f"and internally we track everything as a '{internal_term}'."
    )

    doc.add_page_break()

    # 1) Development Path
    _add_heading(doc, "1) Development Path (Roadmap to Launch)", level=1)
    _add_heading(doc, "Product scope (from the plan)", level=2)
    _add_kv_table(
        doc,
        [
            ("Launch target", "2026 Diwali"),
            ("Customer types", "Home owners, hotel owners, PG/hostel, rented homes"),
            (
                "Service categories",
                "AC, Washing Machine, TV, Fridge, Geyser, RO Filter, Cooler, Speaker/Multimedia",
            ),
            (
                "Core system needs",
                "Booking, technician assignment, work order tracking, diagnosis → estimate → approval, "
                "photos (before/after), signature/OTP confirmation, billing, ratings, admin dashboard",
            ),
        ],
    )

    _add_heading(doc, "Phase 0 — Product definition (1–2 weeks)", level=2)
    _add_bullets(
        doc,
        [
            "Finalize service catalog (categories → sub-issues → standard check-up/repair/service types).",
            "Define pricing rules: check-up/visit fee, platform fee, parts pricing rules, cancellation/refund.",
            "Define roles & permissions: Customer, Technician, Admin/Operations, Support (optional).",
            "Pick MVP geography (start with one city/area).",
        ],
    )

    _add_heading(doc, "Phase 1 — UX + data + architecture (2–3 weeks)", level=2)
    _add_bullets(
        doc,
        [
            "Finalize user journeys for Customer, Technician, Admin (see section 2).",
            "Define booking/work-order status machine + SLAs (accept/arrive/complete).",
            "Design core data models (users, technicians, bookings/work orders, estimates, payments, reviews).",
            "Define API contract (endpoints + request/response).",
            "Non-functional: backups, audit logs, rate limits, security, privacy.",
        ],
    )

    _add_heading(doc, "Phase 2 — MVP build (8–12 weeks)", level=2)
    _add_bullets(
        doc,
        [
            "Customer portal: service selection, booking, status tracking, estimate approval, payment, history, rating, support.",
            "Technician portal: profile/KYC, availability, job accept/decline, navigation, diagnosis, estimate builder, approval capture, photos, completion, earnings.",
            "Admin portal: service catalog & pricing, technician verification, live job monitoring, disputes/refunds, analytics.",
        ],
    )

    _add_heading(doc, "Phase 3 — QA + pilot (3–5 weeks)", level=2)
    _add_bullets(
        doc,
        [
            "UAT with real technicians and customers (target: 50–200 bookings).",
            "Operational SOPs for support, escalations, cancellations/refunds.",
            "Abuse controls: OTP throttling, cancellation rules, audit trails for estimate approvals.",
        ],
    )

    _add_heading(doc, "Phase 4 — Launch + scale (ongoing)", level=2)
    _add_bullets(
        doc,
        [
            "City expansion playbook (service area definition, onboarding, local support).",
            "Quality controls: rating thresholds, rework rules, technician penalties.",
            "Growth loops: referrals, service landing pages by city, partnerships.",
        ],
    )

    doc.add_page_break()

    # 2) Workflows
    _add_heading(doc, "2) Work Flow of the App (Customer + Technician + Admin)", level=1)

    _add_heading(doc, "Customer workflow (happy path)", level=2)
    _add_numbered(
        doc,
        [
            "Sign up / login (OTP).",
            "Select service category and issue (e.g., AC → No cooling / Service / Gas leak).",
            "Enter details: address, preferred slot, issue notes, optional photos.",
            "Price preview: check-up/visit + platform fee; parts extra after diagnosis.",
            f"{customer_term} created.",
            "Technician assigned (matching by service, distance, rating, availability).",
            "Tracking: technician status + ETA; chat/call available.",
            "Diagnosis & estimate shared to customer (itemized).",
            "Customer approves estimate (in-app approval + OTP/signature).",
            "Repair execution and progress updates.",
            "Completion: before/after photos, final invoice generated.",
            "Payment confirmation.",
            "Rating + feedback; raise complaint if required.",
        ],
    )

    _add_heading(doc, "Technician workflow (implementation-ready)", level=2)
    _add_numbered(
        doc,
        [
            f"Login → view new/assigned {technician_term.lower()}s.",
            "Accept/decline with timer (decline reason required).",
            "Update status: On route → Arrived.",
            "Diagnosis: select device category, select issues (multi-select), add findings + notes + photos.",
            "Estimate: add parts (name, qty, price), labor/service charges, taxes/platform fees (as per rules).",
            "Capture customer approval (approve button + OTP/signature).",
            "Repair log: key steps and before/after photos (timestamped).",
            f"Final testing checklist; mark {technician_term.lower()} completed.",
            "Payment status confirmed; earnings updated.",
        ],
    )

    _add_heading(doc, "Admin workflow", level=2)
    _add_bullets(
        doc,
        [
            "Technician onboarding and verification (documents, services, service area, background verification flag).",
            "Live booking monitoring with SLA timers (assigned/accepted/arrived/completed).",
            "Service catalog and pricing management.",
            "Dispute and complaint handling: refunds/partial refunds, penalties, suspensions.",
            "Reports: completion rate, cancellations, repeat issues, technician performance, revenue.",
        ],
    )

    _add_heading(doc, f"Recommended {internal_term.lower()} statuses", level=2)
    _add_bullets(
        doc,
        [
            "draft → created → assigned → accepted → on_route → arrived",
            "diagnosis_in_progress → estimate_sent → approved → repair_in_progress",
            "completed → paid → closed",
            "Side paths: cancelled_by_customer, cancelled_by_tech, no_show, disputed, refunded",
        ],
    )

    doc.add_page_break()

    # 3) App structure
    _add_heading(doc, "3) App Structure (Pages, Modules, APIs, Data)", level=1)

    _add_heading(doc, "3.1 Web apps (3 portals)", level=2)
    _add_bullets(
        doc,
        [
            "Customer portal: app.servesetu.com",
            "Technician portal: tech.servesetu.com",
            "Admin portal: admin.servesetu.com",
        ],
    )

    _add_heading(doc, "3.2 Pages (minimum set)", level=2)
    _add_heading(doc, "Customer pages", level=3)
    _add_bullets(
        doc,
        [
            "Auth (OTP login), Profile",
            "Home (service categories)",
            "Service details (issue selection + pricing rules)",
            "Booking (address, slot, notes, upload images)",
            "Tracking (status timeline + ETA + contact/chat)",
            "Estimate approval (itemized estimate, approve/reject)",
            "Payment (invoice + pay now)",
            "History (past bookings, invoices)",
            "Rating & Support (review, raise complaint)",
        ],
    )

    _add_heading(doc, "Technician pages", level=3)
    _add_bullets(
        doc,
        [
            "Auth",
            "KYC/Profile (services, area, availability, documents)",
            "Work Orders list (new + active + completed)",
            "Work Order details (customer info, location, notes)",
            "Diagnosis (device, issues, findings)",
            "Estimate builder (parts + labor + fees)",
            "Approval capture (signature/OTP)",
            "Repair log (steps + photos)",
            "Earnings (daily/weekly, payouts)",
        ],
    )

    _add_heading(doc, "Admin pages", level=3)
    _add_bullets(
        doc,
        [
            "Dashboard",
            "Bookings monitor",
            "Technicians management (verify/suspend)",
            "Customers management",
            "Service catalog & pricing",
            "Disputes/refunds",
            "Reports/analytics",
        ],
    )

    _add_heading(doc, "3.3 Core modules (backend + shared)", level=2)
    _add_bullets(
        doc,
        [
            "Auth & identity (OTP, session/JWT, roles)",
            "User profiles (customer/technician/admin)",
            "Address & service area",
            "Service catalog (categories, issues, checklists)",
            "Booking (scheduling, assignment, status machine)",
            "Technician availability",
            "Estimate & approval (itemized quote + approval audit log)",
            "Payments (gateway, webhooks, refunds)",
            "Media storage (before/after photos + timestamps)",
            "Ratings & reviews",
            "Notifications (SMS/WhatsApp/email/push; phase-based)",
            "Support & disputes (tickets, refunds, penalties)",
            "Audit logging (status/estimate changes with actor + timestamp)",
        ],
    )

    _add_heading(doc, "3.4 Suggested API list (MVP)", level=2)
    _add_heading(doc, "Auth", level=3)
    _add_bullets(doc, ["POST /auth/request-otp", "POST /auth/verify-otp", "POST /auth/logout"])

    _add_heading(doc, "Catalog + pricing", level=3)
    _add_bullets(
        doc,
        [
            "GET /services/categories",
            "GET /services/categories/:id/issues",
            "GET /pricing/preview (inputs: category, issue, location)",
        ],
    )

    _add_heading(doc, "Bookings (customer)", level=3)
    _add_bullets(doc, ["POST /bookings", "GET /bookings/:id", "GET /bookings (history)", "POST /bookings/:id/cancel"])

    _add_heading(doc, "Work Orders (technician)", level=3)
    _add_bullets(
        doc,
        [
            "GET /tech/work-orders (filters: new/active/completed)",
            "POST /tech/work-orders/:id/accept",
            "POST /tech/work-orders/:id/decline",
            "POST /tech/work-orders/:id/status (on_route/arrived/etc.)",
        ],
    )

    _add_heading(doc, "Diagnosis + estimate + approval", level=3)
    _add_bullets(
        doc,
        [
            "POST /service-requests/:id/diagnosis",
            "POST /service-requests/:id/estimate",
            "POST /service-requests/:id/estimate/approve",
            "POST /service-requests/:id/estimate/reject",
        ],
    )

    _add_heading(doc, "Media", level=3)
    _add_bullets(doc, ["POST /uploads (get upload URL)", "POST /service-requests/:id/photos (before/after refs)"])

    _add_heading(doc, "Payments", level=3)
    _add_bullets(
        doc,
        [
            "POST /payments/create-intent",
            "POST /payments/webhook (gateway callback)",
            "POST /payments/:id/refund (admin)",
        ],
    )

    _add_heading(doc, "Reviews", level=3)
    _add_bullets(doc, ["POST /service-requests/:id/review", "GET /tech/:id/reviews"])

    _add_heading(doc, "Admin", level=3)
    _add_bullets(
        doc,
        [
            "GET /admin/bookings",
            "PATCH /admin/technicians/:id (verify/suspend)",
            "PATCH /admin/catalog/*",
            "POST /admin/disputes/:id/resolve",
        ],
    )

    _add_heading(doc, "3.5 Data entities (minimum)", level=2)
    _add_bullets(
        doc,
        [
            "User(id, role, phone, name, status)",
            "CustomerProfile(userId, addresses[])",
            "TechnicianProfile(userId, services[], serviceArea, KYCStatus, ratingAvg)",
            "ServiceCategory(id, name)",
            "ServiceIssueTemplate(id, categoryId, name)",
            "ServiceRequest(id, customerId, technicianId, categoryId, issueIds, address, slot, status, timestamps)",
            "Diagnosis(serviceRequestId, findings, notes, media[])",
            "Estimate(serviceRequestId, lineItems[], totals, approvedAt, approvedBy)",
            "Payment(serviceRequestId, amount, status, gatewayRef)",
            "Review(serviceRequestId, rating, comment)",
            "AuditLog(entityType, entityId, action, actorId, before/after, ts)",
        ],
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    output_path = repo_root / "docs" / "ServeSetu_Product_Blueprint.docx"
    build_doc(output_path)
    print(f"Wrote: {output_path}")


if __name__ == "__main__":
    main()

