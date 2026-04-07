from __future__ import annotations

from datetime import date
from pathlib import Path


def _h(doc, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def _p(doc, text: str) -> None:
    doc.add_paragraph(text)


def _bullets(doc, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _numbered(doc, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def _table(doc, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build_doc(output_path: Path) -> None:
    from docx import Document

    doc = Document()

    # Terms (user-friendly)
    customer_term = "Booking"
    technician_term = "Work Order"
    internal_term = "Service Request"

    # Title
    doc.add_heading("ServeSetu — Team Execution Plan (Web-first MVP)", level=0)
    _p(doc, f"Generated on: {date.today().isoformat()}")
    _p(
        doc,
        "Goal: ship a web-first MVP (Customer + Technician + Admin portals) quickly, with clean ownership, "
        "weekly milestones, and minimal rework.",
    )
    _p(
        doc,
        f"Terminology: Customer sees '{customer_term}', Technician sees '{technician_term}', "
        f"internally tracked as '{internal_term}'.",
    )

    doc.add_page_break()

    # 1) Team & ownership
    _h(doc, "1) Team roles & ownership", 1)
    _h(doc, "Roles", 2)
    _bullets(
        doc,
        [
            "Dev Team (2 people): build backend + 3 web portals + deployment + QA fixes.",
            "Docs/Content/Ops: policies, agreements, pricing/cancellation rules, website content, support SOPs.",
            "Design: logo + brand kit (colors/fonts) + UI kit components (optional but helpful).",
        ],
    )

    _h(doc, "RACI (who does what)", 2)
    _table(
        doc,
        headers=["Area", "Responsible", "Accountable", "Consulted", "Informed"],
        rows=[
            ["MVP scope + city + services", "Founder/Dev", "Founder", "Ops, Design", "All"],
            ["Pricing rules + refund policy", "Ops", "Founder", "Dev", "All"],
            ["UX flows + screens (Figma)", "Design", "Founder", "Dev, Ops", "All"],
            ["Backend architecture + DB", "Dev", "Founder/Dev lead", "Ops (rules)", "All"],
            ["Customer portal", "Dev", "Dev lead", "Design", "All"],
            ["Technician portal", "Dev", "Dev lead", "Design", "All"],
            ["Admin portal", "Dev", "Founder/Dev lead", "Ops", "All"],
            ["Legal docs (T&C, privacy, agreements)", "Ops", "Founder", "—", "All"],
            ["Deployment + monitoring", "Dev", "Dev lead", "—", "All"],
            ["Pilot operations + support SOPs", "Ops", "Founder", "Dev", "All"],
        ],
    )

    doc.add_page_break()

    # 2) MVP definition (what we will ship)
    _h(doc, "2) Web-first MVP definition (must-have)", 1)
    _h(doc, "MVP success criteria", 2)
    _bullets(
        doc,
        [
            f"Customer can create a {customer_term.lower()} and track status end-to-end.",
            f"Technician can accept a {technician_term.lower()}, update status, diagnose, create estimate, and complete service.",
            "Customer can approve estimate and pay; invoice is generated; review is captured.",
            "Admin can monitor all bookings/work orders, manage technicians, and handle cancellations/refunds.",
        ],
    )

    _h(doc, "MVP scope boundaries (to ship faster)", 2)
    _bullets(
        doc,
        [
            "Start with 2–3 service categories only (expand after pilot).",
            "One city/area first.",
            "Keep matching simple (service + distance + availability).",
            "Keep chat optional (can start with call/WhatsApp link).",
        ],
    )

    _h(doc, "Core status flow (single source of truth)", 2)
    _bullets(
        doc,
        [
            "draft → created → assigned → accepted → on_route → arrived",
            "diagnosis_in_progress → estimate_sent → approved → repair_in_progress",
            "completed → paid → closed",
            "Side paths: cancelled_by_customer, cancelled_by_tech, no_show, disputed, refunded",
        ],
    )

    doc.add_page_break()

    # 3) 12-week execution plan
    _h(doc, "3) 12-week execution plan (weekly milestones)", 1)
    _p(
        doc,
        "Principle: deliver a working demo every week. Do not build 'all features' first—build the end-to-end flow early.",
    )

    _h(doc, "Week 1 — Scope freeze + architecture", 2)
    _bullets(
        doc,
        [
            "Lock MVP city + 2–3 categories + pricing rules (visit fee, platform fee, parts extra).",
            "Finalize booking/work-order statuses + SLA timers.",
            "DB schema v1 + API contract v1.",
            "Wireframes / quick Figma for core pages.",
        ],
    )

    _h(doc, "Week 2 — Project setup + auth + catalog", 2)
    _bullets(
        doc,
        [
            "Repo structure + environments (dev/staging).",
            "OTP login + roles (Customer/Technician/Admin).",
            "Service catalog CRUD (admin) + read (customer).",
        ],
    )

    _h(doc, "Week 3 — Booking creation (customer) + assignment", 2)
    _bullets(
        doc,
        [
            "Customer: create booking (address, slot, issue, notes).",
            "Backend: assignment rules (simple matching).",
            "Admin: view all service requests.",
        ],
    )

    _h(doc, "Week 4 — Technician work orders + status tracking", 2)
    _bullets(
        doc,
        [
            "Technician: work order list + accept/decline.",
            "Technician: status updates (on_route, arrived).",
            "Customer: timeline tracking page.",
        ],
    )

    _h(doc, "Week 5 — Diagnosis + estimate v1", 2)
    _bullets(
        doc,
        [
            "Technician: diagnosis form (device, issues, notes).",
            "Technician: estimate builder (parts + labor + fees).",
            "Customer: view estimate.",
        ],
    )

    _h(doc, "Week 6 — Customer approval + audit log", 2)
    _bullets(
        doc,
        [
            "Customer: approve/reject estimate (OTP/signature).",
            "Audit log: record who changed what and when (estimate + status).",
        ],
    )

    _h(doc, "Week 7 — Media (before/after) + completion flow", 2)
    _bullets(
        doc,
        [
            "Uploads: before/after photos for proof (timestamped).",
            "Technician: completion checklist; mark completed.",
            "Customer: completion confirmation screen.",
        ],
    )

    _h(doc, "Week 8 — Payments + invoice + webhooks", 2)
    _bullets(
        doc,
        [
            "Payment gateway integration (UPI/cards as chosen).",
            "Invoice generation + payment status sync via webhooks.",
            "Admin: refund action (manual first).",
        ],
    )

    _h(doc, "Week 9 — Reviews + basic complaints", 2)
    _bullets(
        doc,
        [
            "Customer: rating & feedback.",
            "Support: complaint ticket (basic) + admin resolution notes.",
        ],
    )

    _h(doc, "Week 10 — Technician verification + admin tools", 2)
    _bullets(
        doc,
        [
            "Technician onboarding fields + document upload.",
            "Admin: verify/suspend technicians.",
            "Admin: pricing/catalog editor.",
        ],
    )

    _h(doc, "Week 11 — QA/UAT + performance hardening", 2)
    _bullets(
        doc,
        [
            "UAT with real technicians and customers.",
            "Fix bugs, polish UX, tighten permissions, rate limits, OTP throttling.",
            "Monitoring/logging + backups enabled.",
        ],
    )

    _h(doc, "Week 12 — Pilot launch + operations", 2)
    _bullets(
        doc,
        [
            "Pilot go-live in one area.",
            "Support SOPs active (refunds, cancellations, dispute handling).",
            "Post-pilot improvements list prepared for next sprint.",
        ],
    )

    doc.add_page_break()

    # 4) Definition of Done (DoD)
    _h(doc, "4) Definition of Done (to keep quality high)", 1)
    _h(doc, "For any feature/module", 2)
    _bullets(
        doc,
        [
            "Works end-to-end on staging environment.",
            "Role-based access checked (customer vs technician vs admin).",
            "Validation + error messages handled.",
            "Audit logs for sensitive actions (status changes, estimate edits, refunds).",
            "Basic analytics events (optional) OR at least server logs for key actions.",
        ],
    )

    _h(doc, "For payments", 2)
    _bullets(
        doc,
        [
            "Webhook verification implemented.",
            "Invoice cannot be marked paid without payment confirmation.",
            "Refund path tested (manual admin action is OK for MVP).",
        ],
    )

    doc.add_page_break()

    # 5) Handoffs & meetings
    _h(doc, "5) Handoffs, meetings, and smooth collaboration", 1)
    _h(doc, "Weekly cadence", 2)
    _bullets(
        doc,
        [
            "1 weekly planning call (30–45 min): confirm this week’s milestone + owners.",
            "2 short async updates/day in your group (morning: plan, evening: progress).",
            "1 weekly demo on staging for the whole team.",
        ],
    )

    _h(doc, "Handoff rules (avoid rework)", 2)
    _bullets(
        doc,
        [
            "Ops finalizes policies (pricing/refund/cancellation) before dev implements logic.",
            "Design provides UI kit + key screens before dev builds the module.",
            "Dev exposes staging links early; Ops tests flows and writes FAQs from real UX.",
        ],
    )

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    output_path = repo_root / "docs" / "ServeSetu_Team_Execution_Plan.docx"
    build_doc(output_path)
    print(f"Wrote: {output_path}")


if __name__ == "__main__":
    main()

